"""
Script para generar el documento Word del Proyecto Final - BistroMaster PRO.
Este script utiliza la librería python-docx para crear un reporte profesional
con el código fuente y la documentación de la API.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# Inicializamos el documento
doc = Document()

# ── Configuración de Márgenes ──
# Establecemos márgenes estándar de 2.5cm en todos los lados.
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Definimos el estilo de fuente base (Normal)
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ── Logotipo de la Institución ──
# Intentamos cargar el logo de la universidad si el archivo existe.
logo_path = os.path.join(os.path.dirname(__file__), "..", "logo_university.png")
if os.path.exists(logo_path):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(1.8))

# ── Encabezado Principal ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Centro de Enlace Universitario Felipe Villanueva")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E) # Azul oscuro profesional

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Materia: Programación III B")
run2.bold = True
run2.font.size = Pt(13)

# ── Información del Estudiante ──
# Usamos una tabla sin bordes para alinear la información personal.
doc.add_paragraph() 

table_info = doc.add_table(rows=3, cols=2)
table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ("Alumno:", "Humberto Alejandro Zambrano"),
    ("Materia:", "Programación III B"),
    ("Tarea:", "Proyecto Final - BistroMaster PRO (API + Frontend)"),
]
for i, (label, value) in enumerate(data):
    cell_l = table_info.cell(i, 0)
    cell_l.text = ""
    run_l = cell_l.paragraphs[0].add_run(label)
    run_l.bold = True
    run_l.font.size = Pt(11)
    
    cell_v = table_info.cell(i, 1)
    cell_v.text = value
    cell_v.paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph() 

# ── Título del Proyecto ──
h = doc.add_heading("BistroMaster PRO: Suite de Gestión de Restaurantes", level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 1. Descripción del Proyecto ──
doc.add_heading("1. Descripción del Proyecto", level=2)
doc.add_paragraph(
    "Este proyecto representa una solución integral de extremo a extremo para la gestión de un restaurante moderno. "
    "Consta de un Backend profesional construido con FastAPI que implementa persistencia de datos relacionales "
    "mediante SQLite, y un Frontend dinámico desarrollado con estándares modernos de diseño web (Glassmorphism), "
    "asegurando una experiencia de usuario fluida y profesional."
)

# ── 2. Arquitectura y Tecnologías ──
doc.add_heading("2. Arquitectura y Tecnologías", level=2)
tech_list = [
    "FastAPI: Framework asíncrono de alto rendimiento para el Backend.",
    "SQLModel / SQLAlchemy: ORM para el manejo de la base de datos relacional.",
    "SQLite: Almacenamiento local persistente.",
    "JWT (JSON Web Tokens): Protocolo de seguridad para sesiones de usuario.",
    "Vanilla JS / CSS Moderno: Frontend ligero, rápido y estéticamente superior.",
]
for t in tech_list:
    doc.add_paragraph(t, style="List Bullet")

# ── 3. Documentación de la API ──
doc.add_heading("3. Documentación de la API", level=2)
endpoints_table = doc.add_table(rows=1, cols=3)
endpoints_table.style = "Light Grid Accent 1"
headers = ["Método", "Ruta", "Funcionalidad"]
for j, h_text in enumerate(headers):
    endpoints_table.cell(0, j).text = h_text

endpoints = [
    ("POST", "/token", "Autenticación de personal (Obtención de JWT)"),
    ("GET", "/menu", "Visualización del catálogo para clientes"),
    ("POST", "/menu", "Gestión de inventario (Añadir nuevos platos)"),
    ("DELETE", "/menu/{id}", "Gestión de inventario (Eliminar platos)"),
    ("POST", "/orders", "Recepción de pedidos en tiempo real"),
    ("GET", "/orders", "Monitor de pedidos para administración"),
    ("PATCH", "/orders/{id}/status", "Control de flujo (Pendiente -> Listo -> etc.)"),
]
for m, r, f in endpoints:
    row = endpoints_table.add_row().cells
    row[0].text = m
    row[1].text = r
    row[2].text = f

# ── 4. Código Fuente por Módulos ──
doc.add_heading("4. Código fuente (Arquitectura Modular)", level=2)

# Lista de archivos que incluiremos en el documento de entrega.
files_to_include = [
    ("main.py", "Gestión de rutas y controladores de la API."),
    ("database.py", "Modelos de la base de datos y configuración del motor."),
    ("schemas.py", "Definición de esquemas de validación Pydantic."),
    ("auth.py", "Seguridad, hashing de contraseñas y lógica de tokens."),
]

for filename, desc in files_to_include:
    doc.add_heading(f"Módulo: {filename}", level=3)
    doc.add_paragraph(desc)
    
    # Intentamos leer el archivo y agregarlo con formato tipográfico de código.
    if os.path.exists(filename):
        with open(filename, "r", encoding="utf-8") as f:
            code = f.read()

        code_table = doc.add_table(rows=1, cols=1)
        cell = code_table.cell(0, 0)
        
        # Aplicamos un sombreado gris claro para que parezca un bloque de código.
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F5F5F5")
        cell._tc.get_or_add_tcPr().append(shading)

        p_code = cell.paragraphs[0]
        for line in code.split("\n"):
            run_code = p_code.add_run(line + "\n")
            run_code.font.name = "Consolas"
            run_code.font.size = Pt(7)

# ── Finalización y Guardado ──
output_filename = "Entrega_Final_BistroMaster_PRO_Zambrano.docx"
output_path = os.path.join(os.path.dirname(__file__), output_filename)
doc.save(output_path)

print(f"✨ ¡Éxito! El reporte de entrega se ha generado en: {output_path}")
