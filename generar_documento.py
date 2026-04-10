"""
Script para generar el documento Word universitario del proyecto final.
Centro de Enlace Universitario Felipe Villanueva
Estudiante: Humberto Alejandro Zambrano
Materia: Programacion III B
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Configurar estilos base ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# ── Leer el codigo fuente de main.py ─────────────────────────────────────────
with open("/home/humberto/source/home_work/proyecto_final/main.py", "r", encoding="utf-8") as f:
    source_code = f.read()


# ── Funciones auxiliares ─────────────────────────────────────────────────────

def add_heading_styled(text, level=1, bold=True, size=None, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    if level == 1:
        run.font.size = Pt(size or 16)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(size or 14)
        run.font.color.rgb = RGBColor(0, 51, 102)
    else:
        run.font.size = Pt(size or 12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    return p


def add_body(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body_no_indent(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = 'Times New Roman'
        run_b.font.size = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_code_block(code_text, font_size=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p


def add_numbered_item(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_n = p.add_run(f"{number}. ")
    run_n.bold = True
    run_n.font.name = 'Times New Roman'
    run_n.font.size = Pt(12)
    run_t = p.add_run(text)
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.63)
    return p


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def format_table_cell(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return cell


def create_table(headers, data, col_widths=None, header_color="003366"):
    rows = len(data) + 1
    cols = len(headers)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        format_table_cell(cell, h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            format_table_cell(cell, val, size=10)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "EBF0FA")

    doc.add_paragraph()
    return table


def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ═══════════════════════════════════════════════════════════════════════════════
#                               PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture("/home/humberto/source/home_work/logo_universidad.png", width=Inches(2.2))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CENTRO DE ENLACE UNIVERSITARIO\nFELIPE VILLANUEVA")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

add_separator()

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROYECTO FINAL")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(18)
p.paragraph_format.space_after = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "API REST para la Gestion de Pedidos de Restaurante\n"
    "con Autenticacion JWT mediante FastAPI"
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

for _ in range(5):
    doc.add_paragraph()

# Tabla de datos del estudiante en la portada
info_table = doc.add_table(rows=3, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Materia:", "Programacion III B"),
    ("Estudiante:", "Humberto Alejandro Zambrano"),
    ("Fecha de entrega:", "Abril de 2026"),
]
for i, (label, value) in enumerate(info_data):
    cell_l = info_table.rows[i].cells[0]
    cell_l.text = ""
    p = cell_l.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    cell_v = info_table.rows[i].cells[1]
    cell_v.text = ""
    p = cell_v.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  {value}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Quitar bordes de la tabla de info
for row in info_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                          INDICE DE CONTENIDOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("INDICE DE CONTENIDOS", level=1, size=18)
add_separator()

indice = [
    ("1.", "Introduccion"),
    ("2.", "Objetivos"),
    ("", "2.1. Objetivo General"),
    ("", "2.2. Objetivos Especificos"),
    ("3.", "Marco Teorico"),
    ("", "3.1. API REST"),
    ("", "3.2. FastAPI"),
    ("", "3.3. JSON Web Tokens (JWT)"),
    ("", "3.4. Pydantic y Validacion de Datos"),
    ("", "3.5. OAuth2 y Seguridad en APIs"),
    ("", "3.6. Cifrado de Contrasenas con Bcrypt"),
    ("4.", "Descripcion del Proyecto"),
    ("", "4.1. Arquitectura del Sistema"),
    ("", "4.2. Modelos de Datos"),
    ("", "4.3. Endpoints de la API"),
    ("", "4.4. Flujo de Autenticacion y Autorizacion"),
    ("5.", "Desarrollo e Implementacion"),
    ("", "5.1. Tecnologias y Dependencias Utilizadas"),
    ("", "5.2. Estructura y Organizacion del Codigo Fuente"),
    ("", "5.3. Mecanismos de Seguridad Implementados"),
    ("", "5.4. Codigo Fuente Completo"),
    ("6.", "Pruebas y Resultados"),
    ("", "6.1. Metodologia de Pruebas"),
    ("", "6.2. Casos de Prueba"),
    ("", "6.3. Resultados Obtenidos"),
    ("7.", "Conclusiones"),
    ("8.", "Recomendaciones"),
    ("9.", "Referencias Bibliograficas"),
    ("10.", "Anexos"),
]

for num, titulo in indice:
    p = doc.add_paragraph()
    if num:
        run_n = p.add_run(f"{num} ")
        run_n.bold = True
        run_n.font.name = 'Times New Roman'
        run_n.font.size = Pt(12)
        run_t = p.add_run(titulo)
        run_t.bold = True
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(12)
    else:
        p.paragraph_format.left_indent = Cm(1.5)
        run_t = p.add_run(titulo)
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                          1. INTRODUCCION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("1. INTRODUCCION", level=1)

add_body(
    "En el contexto actual de la transformacion digital, los sistemas de informacion "
    "desempenan un papel fundamental en la optimizacion de procesos operativos en "
    "diversos sectores de la economia. La industria de la restauracion no es ajena "
    "a esta tendencia: la necesidad de gestionar pedidos de manera eficiente, precisa "
    "y segura se ha convertido en un requisito indispensable para garantizar la "
    "satisfaccion del cliente y la competitividad del negocio."
)
add_body(
    "Las Interfaces de Programacion de Aplicaciones (APIs) de tipo REST "
    "(Representational State Transfer) se han consolidado como el estandar de facto "
    "para la comunicacion entre sistemas distribuidos en la web moderna. Su adopcion "
    "permite desacoplar la logica del servidor de la interfaz de usuario, facilitando "
    "la integracion con multiples plataformas (aplicaciones moviles, sistemas de punto "
    "de venta, paneles administrativos, entre otros)."
)
add_body(
    "El presente trabajo documenta el desarrollo de una API REST para la gestion "
    "integral de pedidos de un restaurante, implementada con el framework FastAPI "
    "del ecosistema Python. El sistema incorpora autenticacion basada en JSON Web "
    "Tokens (JWT), cifrado de contrasenas mediante el algoritmo bcrypt, y validacion "
    "rigurosa de datos a traves de modelos Pydantic. Estas tecnologias, en conjunto, "
    "permiten construir un servicio web robusto, seguro y de alto rendimiento."
)
add_body(
    "A lo largo de este documento se presentan los fundamentos teoricos que sustentan "
    "las decisiones tecnologicas, la descripcion detallada de la arquitectura y los "
    "componentes del sistema, el codigo fuente completo con su respectivo analisis, "
    "las pruebas realizadas y los resultados obtenidos, asi como las conclusiones "
    "y recomendaciones derivadas del proceso de desarrollo."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                          2. OBJETIVOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("2. OBJETIVOS", level=1)

add_heading_styled("2.1. Objetivo General", level=2)
add_body(
    "Disenar, desarrollar y documentar una API REST funcional para la gestion de "
    "pedidos de un restaurante, empleando FastAPI como framework de desarrollo y "
    "JSON Web Tokens (JWT) como mecanismo de autenticacion y autorizacion, aplicando "
    "buenas practicas de ingenieria de software y seguridad informatica."
)

add_heading_styled("2.2. Objetivos Especificos", level=2)
objectives = [
    "Disenar modelos de datos que representen adecuadamente las entidades del dominio "
    "(menu, pedidos, usuarios) mediante esquemas Pydantic con validaciones de integridad.",
    "Implementar operaciones CRUD (Create, Read, Update, Delete) completas para la "
    "gestion del menu y los pedidos del restaurante.",
    "Integrar un flujo de autenticacion OAuth2 con Password Flow, generando tokens "
    "JWT firmados con algoritmo HS256 y tiempo de expiracion configurable.",
    "Aplicar cifrado bcrypt con salt automatico para el almacenamiento seguro de "
    "contrasenas de usuarios.",
    "Garantizar la proteccion de endpoints administrativos mediante inyeccion de "
    "dependencias de autenticacion en FastAPI.",
    "Validar automaticamente la estructura y los tipos de datos de todas las "
    "peticiones entrantes mediante los mecanismos nativos de Pydantic.",
    "Generar documentacion interactiva de la API de forma automatica a traves de "
    "la interfaz Swagger UI proporcionada por FastAPI.",
]
for i, obj in enumerate(objectives, 1):
    add_numbered_item(i, obj)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                          3. MARCO TEORICO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("3. MARCO TEORICO", level=1)

add_heading_styled("3.1. API REST", level=2)
add_body(
    "Una API (Application Programming Interface) es un conjunto de definiciones y "
    "protocolos que permiten la comunicacion entre componentes de software. El estilo "
    "arquitectonico REST (Representational State Transfer), propuesto por Roy Fielding "
    "en su disertacion doctoral en el ano 2000, establece un conjunto de restricciones "
    "para el diseno de sistemas distribuidos basados en el protocolo HTTP."
)
add_body(
    "Los principios fundamentales de REST incluyen: (a) la arquitectura cliente-servidor, "
    "que separa las responsabilidades de la interfaz de usuario y el almacenamiento de "
    "datos; (b) la comunicacion sin estado (stateless), donde cada peticion contiene toda "
    "la informacion necesaria para ser procesada; (c) la identificacion unica de recursos "
    "mediante URIs (Uniform Resource Identifiers); y (d) el uso de los metodos estandar "
    "de HTTP (GET, POST, PUT, PATCH, DELETE) para realizar operaciones sobre dichos recursos."
)
add_body(
    "Las APIs REST utilizan comunmente el formato JSON (JavaScript Object Notation) "
    "para la representacion de datos, debido a su ligereza, legibilidad y amplia "
    "compatibilidad con lenguajes de programacion modernos."
)

add_heading_styled("3.2. FastAPI", level=2)
add_body(
    "FastAPI es un framework web moderno para la construccion de APIs con Python, "
    "creado por Sebastian Ramirez. Se basa en los type hints (anotaciones de tipo) "
    "estandar de Python 3.7+ y utiliza internamente dos componentes clave: Starlette "
    "para el manejo de peticiones HTTP de forma asincrona, y Pydantic para la "
    "validacion y serializacion de datos."
)
add_body(
    "Entre las caracteristicas mas relevantes de FastAPI se encuentran:"
)
add_bullet("Alto rendimiento, comparable a frameworks como Node.js (Express) y Go (Gin), gracias a su naturaleza asincrona basada en ASGI.", bold_prefix="Rendimiento: ")
add_bullet("Generacion automatica de documentacion interactiva en los formatos Swagger UI y ReDoc, sin necesidad de configuracion adicional.", bold_prefix="Documentacion automatica: ")
add_bullet("Validacion automatica de tipos de datos, longitudes, rangos y patrones, derivada directamente de las anotaciones de tipo de Python.", bold_prefix="Validacion de datos: ")
add_bullet("Sistema de inyeccion de dependencias integrado que facilita la reutilizacion de logica comun, como la autenticacion.", bold_prefix="Inyeccion de dependencias: ")

add_heading_styled("3.3. JSON Web Tokens (JWT)", level=2)
add_body(
    "JSON Web Token (JWT), definido en el RFC 7519 de la Internet Engineering Task "
    "Force (IETF), es un estandar abierto que establece un formato compacto y "
    "autocontenido para la transmision segura de informacion entre partes como un "
    "objeto JSON. La informacion contenida en un JWT puede ser verificada y confiable "
    "gracias a su firma digital."
)
add_body("La estructura de un JWT consta de tres componentes separados por puntos:")
add_numbered_item(1,
    "Header: contiene el tipo de token (JWT) y el algoritmo de firma utilizado "
    "(por ejemplo, HS256)."
)
add_numbered_item(2,
    "Payload: contiene los claims, que son declaraciones sobre el usuario y "
    "metadatos adicionales (como el tiempo de expiracion)."
)
add_numbered_item(3,
    "Signature: se genera aplicando el algoritmo especificado en el header "
    "sobre el header codificado, el payload codificado y una clave secreta. "
    "Esta firma garantiza la integridad del token."
)
add_body(
    "En el contexto de APIs web, el flujo tipico consiste en: el cliente se "
    "autentica enviando sus credenciales al servidor; el servidor genera un JWT "
    "y lo retorna al cliente; el cliente incluye este token en el encabezado "
    "Authorization de las peticiones subsiguientes, permitiendo al servidor "
    "verificar la identidad del solicitante sin necesidad de mantener sesiones "
    "en el servidor (stateless authentication)."
)

add_heading_styled("3.4. Pydantic y Validacion de Datos", level=2)
add_body(
    "Pydantic es una libreria de Python que proporciona validacion de datos y "
    "gestion de configuraciones mediante las anotaciones de tipo nativas del "
    "lenguaje. Los modelos Pydantic (clases que heredan de BaseModel) definen "
    "la estructura esperada de los datos, incluyendo tipos, restricciones de "
    "longitud, rangos numericos y patrones de expresiones regulares."
)
add_body(
    "Cuando FastAPI recibe una peticion, los datos del cuerpo (body), los "
    "parametros de ruta (path parameters) y los parametros de consulta (query "
    "parameters) se validan automaticamente contra los modelos Pydantic definidos. "
    "Si los datos no cumplen con las restricciones, FastAPI retorna una respuesta "
    "HTTP 422 (Unprocessable Entity) con un detalle preciso del error de validacion, "
    "sin necesidad de escribir codigo de validacion manual."
)

add_heading_styled("3.5. OAuth2 y Seguridad en APIs", level=2)
add_body(
    "OAuth2, definido en el RFC 6749, es un framework de autorizacion que permite "
    "a las aplicaciones obtener acceso limitado a recursos protegidos. Existen "
    "multiples flujos (grant types) definidos en el estandar; el utilizado en este "
    "proyecto es el Resource Owner Password Credentials Grant (Password Flow), en "
    "el cual el usuario proporciona directamente su nombre de usuario y contrasena "
    "a la aplicacion para obtener un token de acceso."
)
add_body(
    "FastAPI proporciona soporte nativo para OAuth2 a traves de la clase "
    "OAuth2PasswordBearer, que actua como un esquema de seguridad que extrae "
    "automaticamente el token del encabezado Authorization de cada peticion. "
    "Este mecanismo se integra con el sistema de inyeccion de dependencias de "
    "FastAPI, permitiendo proteger endpoints de manera declarativa."
)

add_heading_styled("3.6. Cifrado de Contrasenas con Bcrypt", level=2)
add_body(
    "Bcrypt es un algoritmo de hashing de contrasenas disenado por Niels Provos "
    "y David Mazieres en 1999, basado en el cifrado Blowfish. A diferencia de "
    "algoritmos de hash genericos como SHA-256, bcrypt incorpora un factor de "
    "trabajo (cost factor) ajustable que controla el numero de iteraciones del "
    "algoritmo, haciendolo deliberadamente lento para dificultar ataques de fuerza "
    "bruta y de diccionario."
)
add_body(
    "Ademas, bcrypt genera automaticamente un salt (valor aleatorio) unico para "
    "cada contrasena, lo que garantiza que dos usuarios con la misma contrasena "
    "tendran hashes diferentes. En el contexto de este proyecto, la libreria "
    "passlib se utiliza como interfaz de alto nivel para bcrypt, proporcionando "
    "funciones para hashear y verificar contrasenas de manera segura."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                     4. DESCRIPCION DEL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("4. DESCRIPCION DEL PROYECTO", level=1)

add_heading_styled("4.1. Arquitectura del Sistema", level=2)
add_body(
    "El sistema sigue una arquitectura monolitica de una sola capa (single-tier), "
    "implementada en un unico archivo Python (main.py) que contiene la totalidad "
    "de la logica de la aplicacion. Esta decision arquitectonica es apropiada para "
    "un proyecto de alcance academico y prototipado rapido, ya que simplifica el "
    "desarrollo, las pruebas y el despliegue."
)
add_body(
    "La aplicacion utiliza almacenamiento en memoria mediante estructuras de datos "
    "nativas de Python (listas y diccionarios) como base de datos simulada. Si bien "
    "este enfoque implica que los datos no persisten entre reinicios del servidor, "
    "permite concentrar el desarrollo en la logica de negocio y los mecanismos de "
    "seguridad sin la complejidad adicional de un sistema de gestion de bases de datos."
)
add_body("El flujo de procesamiento de una peticion sigue la siguiente secuencia:")
add_numbered_item(1, "El cliente (navegador web, Postman, curl, aplicacion movil) envia una peticion HTTP al servidor.")
add_numbered_item(2, "El servidor ASGI (Uvicorn) recibe la peticion y la enruta al handler correspondiente de FastAPI.")
add_numbered_item(3, "FastAPI valida automaticamente los datos de entrada contra los modelos Pydantic definidos.")
add_numbered_item(4, "Para endpoints protegidos, el sistema de dependencias ejecuta la funcion get_current_user, que extrae y verifica el token JWT del encabezado Authorization.")
add_numbered_item(5, "Si la autenticacion y validacion son exitosas, se ejecuta la logica de negocio del endpoint.")
add_numbered_item(6, "La respuesta se serializa a formato JSON y se retorna al cliente con el codigo de estado HTTP apropiado.")

add_heading_styled("4.2. Modelos de Datos", level=2)
add_body(
    "El sistema define seis modelos Pydantic que representan las entidades del "
    "dominio y los objetos de transferencia de datos. A continuacion se detallan "
    "los modelos principales:"
)

add_body("Tabla 1. Modelo MenuItem (Item del Menu)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
create_table(
    headers=["Campo", "Tipo de Dato", "Restricciones", "Descripcion"],
    data=[
        ("id", "int", "Requerido", "Identificador unico del item"),
        ("name", "str", "max_length=100", "Nombre del platillo o bebida"),
        ("description", "str | None", "max_length=255", "Descripcion del item (opcional)"),
        ("price", "float", "gt=0 (mayor que cero)", "Precio unitario del item"),
        ("category", "str", "max_length=50", "Categoria: Entrada, Fuerte, Postre, Bebida"),
    ]
)

add_body("Tabla 2. Modelo Order (Pedido)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
create_table(
    headers=["Campo", "Tipo de Dato", "Restricciones", "Descripcion"],
    data=[
        ("id", "int", "Requerido", "Identificador unico del pedido"),
        ("customer_name", "str", "max_length=100", "Nombre del cliente que realiza el pedido"),
        ("items", "List[OrderItem]", "Requerido", "Lista de items con menu_item_id y quantity"),
        ("total_price", "float", "Default: 0.0", "Precio total (calculado por el servidor)"),
        ("status", "str", "Patron regex validado", "pendiente | en_preparacion | listo | entregado | cancelado"),
        ("created_at", "datetime", "Auto (UTC)", "Fecha y hora de creacion del pedido"),
    ]
)

add_body("Tabla 3. Modelo OrderItem (Item de un Pedido)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
create_table(
    headers=["Campo", "Tipo de Dato", "Restricciones", "Descripcion"],
    data=[
        ("menu_item_id", "int", "Requerido", "ID del item del menu solicitado"),
        ("quantity", "int", "ge=1 (minimo 1)", "Cantidad solicitada del item"),
    ]
)

add_body("Tabla 4. Modelos de Autenticacion", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
create_table(
    headers=["Modelo", "Campos", "Proposito"],
    data=[
        ("User", "username, password", "Representacion de credenciales de usuario"),
        ("Token", "access_token, token_type", "Respuesta del endpoint de autenticacion"),
        ("TokenData", "username (opcional)", "Datos extraidos del payload del JWT"),
    ]
)

add_heading_styled("4.3. Endpoints de la API", level=2)
add_body(
    "La API expone nueve endpoints organizados en tres grupos funcionales: "
    "autenticacion, gestion de menu y gestion de pedidos. La siguiente tabla "
    "detalla cada uno de ellos:"
)

add_body("Tabla 5. Endpoints de la API REST", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
create_table(
    headers=["Metodo HTTP", "Ruta", "Descripcion", "Autenticacion", "Codigo Exito"],
    data=[
        ("POST", "/token", "Obtener token de acceso JWT", "No", "200"),
        ("GET", "/menu", "Listar todos los items del menu", "No", "200"),
        ("GET", "/menu/{id}", "Obtener un item del menu por ID", "No", "200"),
        ("POST", "/menu", "Agregar un nuevo item al menu", "Si (JWT)", "201"),
        ("GET", "/orders", "Listar todos los pedidos", "Si (JWT)", "200"),
        ("POST", "/orders", "Crear un nuevo pedido", "No", "201"),
        ("GET", "/orders/{id}", "Obtener un pedido por ID", "Si (JWT)", "200"),
        ("PATCH", "/orders/{id}/status", "Actualizar estado de un pedido", "Si (JWT)", "200"),
        ("DELETE", "/orders/{id}", "Eliminar un pedido", "Si (JWT)", "204"),
    ]
)

add_heading_styled("4.4. Flujo de Autenticacion y Autorizacion", level=2)
add_body(
    "El sistema implementa un flujo de autenticacion basado en OAuth2 Password "
    "Flow con tokens JWT. El proceso completo se describe a continuacion:"
)
add_numbered_item(1,
    "Solicitud de token: El cliente envia una peticion POST al endpoint /token "
    "con las credenciales del usuario (username y password) en formato "
    "application/x-www-form-urlencoded."
)
add_numbered_item(2,
    "Verificacion de credenciales: El servidor busca el usuario en la base de "
    "datos y utiliza bcrypt para comparar la contrasena proporcionada con el "
    "hash almacenado. Si las credenciales son invalidas, se retorna un error "
    "HTTP 401 (Unauthorized)."
)
add_numbered_item(3,
    "Generacion del token: Si la autenticacion es exitosa, se genera un token "
    "JWT que contiene el nombre de usuario en el claim 'sub' (subject) y un "
    "tiempo de expiracion de 30 minutos. El token se firma con la clave secreta "
    "del servidor utilizando el algoritmo HS256."
)
add_numbered_item(4,
    "Uso del token: El cliente almacena el token y lo incluye en el encabezado "
    "HTTP 'Authorization: Bearer <token>' en cada peticion a un endpoint protegido."
)
add_numbered_item(5,
    "Validacion en cada peticion: La dependencia get_current_user, inyectada "
    "automaticamente por FastAPI, decodifica el token, verifica su firma e "
    "integridad, comprueba que no haya expirado, y extrae la identidad del "
    "usuario. Si alguna verificacion falla, se retorna un error HTTP 401."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                  5. DESARROLLO E IMPLEMENTACION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("5. DESARROLLO E IMPLEMENTACION", level=1)

add_heading_styled("5.1. Tecnologias y Dependencias Utilizadas", level=2)
add_body(
    "El desarrollo del proyecto se realizo utilizando el lenguaje de programacion "
    "Python en su version 3.12. Las dependencias del proyecto se gestionan mediante "
    "un archivo requirements.txt y se detallan en la siguiente tabla:"
)

add_body("Tabla 6. Dependencias del Proyecto", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
create_table(
    headers=["Dependencia", "Version", "Proposito"],
    data=[
        ("fastapi", "Ultima estable", "Framework principal para la construccion de la API REST"),
        ("uvicorn", "Ultima estable", "Servidor ASGI de alto rendimiento para ejecutar la aplicacion"),
        ("pydantic", "v2.x", "Validacion de datos y definicion de esquemas de entrada/salida"),
        ("python-jose[cryptography]", "Ultima estable", "Generacion y verificacion de tokens JWT"),
        ("passlib[bcrypt]", "Ultima estable", "Cifrado seguro de contrasenas con algoritmo bcrypt"),
        ("python-multipart", "Ultima estable", "Procesamiento de formularios (OAuth2PasswordRequestForm)"),
    ]
)

add_heading_styled("5.2. Estructura y Organizacion del Codigo Fuente", level=2)
add_body(
    "El codigo fuente se organiza en un unico archivo main.py estructurado en "
    "secciones logicas claramente delimitadas. Esta organizacion facilita la "
    "lectura, el mantenimiento y la comprension del flujo de la aplicacion:"
)

add_body("Tabla 7. Organizacion del Codigo Fuente", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
create_table(
    headers=["Seccion", "Lineas", "Contenido"],
    data=[
        ("Importaciones", "1 - 8", "Librerias externas: FastAPI, JWT, Pydantic, passlib"),
        ("Configuracion JWT", "10 - 12", "Clave secreta, algoritmo HS256, expiracion 30 min"),
        ("Configuracion de Seguridad", "14 - 16", "Contexto bcrypt y esquema OAuth2PasswordBearer"),
        ("Modelos Pydantic", "18 - 48", "MenuItem, OrderItem, Order, User, Token, TokenData"),
        ("Base de Datos en Memoria", "50 - 65", "Listas/diccionarios: menu_db, orders_db, users_db"),
        ("Funciones Auxiliares", "67 - 80", "verify_password(), create_access_token()"),
        ("Dependencia de Autenticacion", "82 - 102", "get_current_user(): validacion de token JWT"),
        ("Instancia de la Aplicacion", "104 - 110", "Creacion de la app FastAPI con metadatos"),
        ("Endpoints de Autenticacion", "112 - 128", "POST /token: login y generacion de token"),
        ("Endpoints del Menu", "130 - 148", "GET/POST /menu: consulta y gestion del menu"),
        ("Endpoints de Pedidos", "150 - 197", "CRUD completo: crear, leer, actualizar, eliminar pedidos"),
    ]
)

add_heading_styled("5.3. Mecanismos de Seguridad Implementados", level=2)
add_body(
    "La seguridad del sistema se implementa en multiples capas, siguiendo el "
    "principio de defensa en profundidad (defense in depth):"
)

add_body("a) Cifrado de contrasenas:", bold=True)
add_body(
    "Las contrasenas de los usuarios nunca se almacenan en texto plano. Se utiliza "
    "el algoritmo bcrypt a traves de la libreria passlib, que genera automaticamente "
    "un salt unico para cada hash. La funcion verify_password() realiza la comparacion "
    "segura entre la contrasena proporcionada y el hash almacenado, evitando ataques "
    "de timing."
)

add_body("b) Tokens JWT con expiracion:", bold=True)
add_body(
    "Cada token generado tiene una vida util de 30 minutos (configurable mediante "
    "la constante ACCESS_TOKEN_EXPIRE_MINUTES). Una vez expirado, el token es rechazado "
    "automaticamente por el sistema de validacion, obligando al usuario a autenticarse "
    "nuevamente. Esto reduce la ventana de exposicion en caso de que un token sea "
    "comprometido."
)

add_body("c) Proteccion de endpoints mediante dependencias:", bold=True)
add_body(
    "Los endpoints que requieren autenticacion utilizan el parametro "
    "current_user: dict = Depends(get_current_user), que es evaluado "
    "automaticamente por FastAPI antes de ejecutar la logica del endpoint. Si "
    "el token es invalido, expirado o ausente, la peticion se rechaza con un "
    "codigo HTTP 401 antes de alcanzar la logica de negocio."
)

add_body("d) Validacion de datos de entrada:", bold=True)
add_body(
    "Los modelos Pydantic aplican restricciones estrictas a todos los datos "
    "de entrada: longitudes maximas para cadenas de texto (max_length), valores "
    "minimos para cantidades (ge=1), valores positivos para precios (gt=0), y "
    "patrones de expresiones regulares para estados validos de pedidos. Cualquier "
    "dato que no cumpla con estas restricciones es rechazado automaticamente con "
    "un error HTTP 422."
)

doc.add_page_break()

add_heading_styled("5.4. Codigo Fuente Completo", level=2)
add_body(
    "A continuacion se presenta el codigo fuente completo del archivo main.py, "
    "que constituye la totalidad de la logica de la aplicacion:"
)

add_body("Archivo: main.py", bold=True, italic=True)

# Dividir el codigo en bloques para evitar un unico bloque gigante
code_sections = source_code.split("\n# ─")
for i, section in enumerate(code_sections):
    if i > 0:
        section = "# ─" + section
    add_code_block(section.rstrip(), font_size=8)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                     6. PRUEBAS Y RESULTADOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("6. PRUEBAS Y RESULTADOS", level=1)

add_heading_styled("6.1. Metodologia de Pruebas", level=2)
add_body(
    "Las pruebas del sistema se realizaron de manera manual utilizando dos "
    "herramientas principales: la interfaz Swagger UI integrada en FastAPI "
    "(accesible en la ruta /docs del servidor) y la herramienta de linea de "
    "comandos curl. La interfaz Swagger UI permite ejecutar peticiones directamente "
    "desde el navegador con una interfaz grafica intuitiva, mientras que curl "
    "proporciona control total sobre los encabezados y el cuerpo de las peticiones."
)
add_body(
    "El servidor de desarrollo se ejecuto mediante el comando:"
)
add_code_block("uvicorn main:app --reload --port 8002")
add_body(
    "El parametro --reload activa la recarga automatica del servidor ante cambios "
    "en el codigo fuente, facilitando el ciclo de desarrollo y pruebas."
)

add_heading_styled("6.2. Casos de Prueba", level=2)

add_body("Caso de prueba 1: Autenticacion exitosa", bold=True)
add_body("Objetivo: Verificar que un usuario con credenciales validas obtiene un token JWT.")
add_code_block(
    'Peticion:\n'
    'POST /token\n'
    'Content-Type: application/x-www-form-urlencoded\n'
    'Body: username=chef&password=cocina2024\n\n'
    'Respuesta esperada (HTTP 200):\n'
    '{\n'
    '    "access_token": "<token_jwt>",\n'
    '    "token_type": "bearer"\n'
    '}'
)

add_body("Caso de prueba 2: Autenticacion fallida", bold=True)
add_body("Objetivo: Verificar que credenciales invalidas son rechazadas.")
add_code_block(
    'Peticion:\n'
    'POST /token\n'
    'Body: username=chef&password=incorrecta\n\n'
    'Respuesta esperada (HTTP 401):\n'
    '{\n'
    '    "detail": "Usuario o contrasena incorrectos"\n'
    '}'
)

add_body("Caso de prueba 3: Consulta publica del menu", bold=True)
add_body("Objetivo: Verificar que cualquier usuario puede consultar el menu sin autenticacion.")
add_code_block(
    'Peticion:\n'
    'GET /menu\n\n'
    'Respuesta esperada (HTTP 200):\n'
    '[\n'
    '    {"id": 1, "name": "Pizza Napolitana", "price": 12.50, ...},\n'
    '    {"id": 2, "name": "Hamburguesa Clasica", "price": 8.90, ...},\n'
    '    {"id": 3, "name": "Cerveza Artesanal", "price": 4.50, ...}\n'
    ']'
)

add_body("Caso de prueba 4: Creacion de pedido con calculo automatico", bold=True)
add_body("Objetivo: Verificar que el sistema calcula correctamente el precio total del pedido.")
add_code_block(
    'Peticion:\n'
    'POST /orders\n'
    'Content-Type: application/json\n'
    'Body:\n'
    '{\n'
    '    "id": 101,\n'
    '    "customer_name": "Juan Perez",\n'
    '    "items": [{"menu_item_id": 1, "quantity": 2}]\n'
    '}\n\n'
    'Respuesta esperada (HTTP 201):\n'
    '{\n'
    '    "id": 101,\n'
    '    "customer_name": "Juan Perez",\n'
    '    "total_price": 25.00,\n'
    '    "status": "pendiente",\n'
    '    ...\n'
    '}'
)

add_body("Caso de prueba 5: Acceso a ruta protegida sin token", bold=True)
add_body("Objetivo: Verificar que las rutas protegidas rechazan peticiones sin autenticacion.")
add_code_block(
    'Peticion:\n'
    'GET /orders\n'
    '(Sin encabezado Authorization)\n\n'
    'Respuesta esperada (HTTP 401):\n'
    '{\n'
    '    "detail": "Not authenticated"\n'
    '}'
)

add_body("Caso de prueba 6: Actualizacion de estado de un pedido", bold=True)
add_body("Objetivo: Verificar que el personal autenticado puede cambiar el estado de un pedido.")
add_code_block(
    'Peticion:\n'
    'PATCH /orders/101/status?status=en_preparacion\n'
    'Authorization: Bearer <token_jwt>\n\n'
    'Respuesta esperada (HTTP 200):\n'
    '{\n'
    '    "id": 101,\n'
    '    "status": "en_preparacion",\n'
    '    ...\n'
    '}'
)

add_body("Caso de prueba 7: Validacion de datos incorrectos", bold=True)
add_body("Objetivo: Verificar que el sistema rechaza datos que no cumplen las restricciones.")
add_code_block(
    'Peticion:\n'
    'POST /orders\n'
    'Body: {"id": 102, "customer_name": "Ana", "items": [{"menu_item_id": 1, "quantity": 0}]}\n\n'
    'Respuesta esperada (HTTP 422):\n'
    'Error de validacion: quantity debe ser >= 1'
)

add_heading_styled("6.3. Resultados Obtenidos", level=2)
add_body(
    "La totalidad de los casos de prueba fueron ejecutados satisfactoriamente. "
    "Los resultados confirman el correcto funcionamiento de los siguientes aspectos "
    "del sistema:"
)
add_bullet("El flujo de autenticacion genera tokens validos y rechaza credenciales incorrectas.")
add_bullet("Los endpoints publicos son accesibles sin autenticacion.")
add_bullet("Los endpoints protegidos rechazan peticiones sin token o con token invalido/expirado.")
add_bullet("El calculo automatico del precio total de los pedidos opera correctamente.")
add_bullet("La validacion de datos rechaza entradas que violan las restricciones definidas en los modelos.")
add_bullet("Las operaciones CRUD (crear, leer, actualizar estado, eliminar) funcionan segun lo esperado.")
add_bullet("La documentacion Swagger UI refleja correctamente todos los endpoints, sus parametros y esquemas de datos.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                          7. CONCLUSIONES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("7. CONCLUSIONES", level=1)

add_body(
    "El desarrollo del presente proyecto permitio alcanzar los siguientes resultados "
    "y aprendizajes:"
)

add_numbered_item(1,
    "Se logro disenar e implementar exitosamente una API REST completa para la "
    "gestion de pedidos de un restaurante, cumpliendo con la totalidad de los "
    "objetivos planteados al inicio del proyecto."
)
add_numbered_item(2,
    "FastAPI demostro ser un framework altamente productivo para el desarrollo "
    "de APIs en Python, destacando su capacidad de generar documentacion "
    "interactiva automaticamente, su sistema de validacion de datos integrado "
    "y su mecanismo de inyeccion de dependencias, que simplifica significativamente "
    "la implementacion de patrones comunes como la autenticacion."
)
add_numbered_item(3,
    "La implementacion de JSON Web Tokens como mecanismo de autenticacion permitio "
    "proteger las operaciones administrativas del sistema de forma efectiva, "
    "manteniendo la naturaleza stateless de la API y evitando la necesidad de "
    "gestionar sesiones en el servidor."
)
add_numbered_item(4,
    "El uso de modelos Pydantic para la validacion de datos proporciona una "
    "primera linea de defensa robusta contra datos malformados o maliciosos, "
    "garantizando la integridad de la informacion desde la capa de entrada "
    "de la aplicacion."
)
add_numbered_item(5,
    "La aplicacion del algoritmo bcrypt para el cifrado de contrasenas sigue "
    "las mejores practicas de seguridad de la industria, asegurando que las "
    "credenciales de los usuarios no puedan ser recuperadas incluso en caso "
    "de una brecha de seguridad en el almacenamiento de datos."
)
add_numbered_item(6,
    "El proyecto sirvio como ejercicio practico integrador de conceptos fundamentales "
    "de programacion: arquitectura cliente-servidor, protocolos HTTP, serializacion "
    "de datos, seguridad informatica, validacion de entradas y diseno de APIs, "
    "consolidando los conocimientos adquiridos en la materia Programacion III B."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                          8. RECOMENDACIONES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("8. RECOMENDACIONES", level=1)

add_body(
    "Con base en la experiencia adquirida durante el desarrollo del proyecto y "
    "con el objetivo de evolucionar el sistema hacia un producto de calidad "
    "profesional, se formulan las siguientes recomendaciones para futuras iteraciones:"
)

add_numbered_item(1,
    "Persistencia de datos: Integrar un sistema de gestion de bases de datos "
    "relacional (PostgreSQL, MySQL) o NoSQL (MongoDB) para garantizar la "
    "persistencia de la informacion entre reinicios del servidor y permitir "
    "consultas complejas."
)
add_numbered_item(2,
    "Sistema de roles y permisos: Implementar un esquema de autorizacion basado "
    "en roles (RBAC - Role-Based Access Control) que permita diferenciar los "
    "privilegios de administradores, cocineros, meseros y clientes."
)
add_numbered_item(3,
    "Refresh tokens: Incorporar un mecanismo de refresh tokens que permita "
    "renovar los tokens de acceso sin requerir que el usuario vuelva a "
    "ingresar sus credenciales, mejorando la experiencia de usuario sin "
    "comprometer la seguridad."
)
add_numbered_item(4,
    "Pruebas automatizadas: Desarrollar un suite de pruebas unitarias y de "
    "integracion utilizando pytest y el cliente de pruebas de FastAPI "
    "(TestClient), garantizando la regresion automatica ante futuros cambios."
)
add_numbered_item(5,
    "Contenerizacion: Empaquetar la aplicacion en un contenedor Docker para "
    "estandarizar el entorno de ejecucion y facilitar el despliegue en "
    "servicios de nube (AWS, GCP, Azure)."
)
add_numbered_item(6,
    "Rate limiting y throttling: Implementar limites de tasa en los endpoints "
    "publicos para prevenir abusos, ataques de denegacion de servicio y "
    "consumo excesivo de recursos."
)
add_numbered_item(7,
    "Logging y monitoreo: Integrar un sistema de logging estructurado "
    "(por ejemplo, con la libreria structlog) y herramientas de monitoreo "
    "para facilitar la depuracion y la observabilidad en entornos de produccion."
)
add_numbered_item(8,
    "Variables de entorno: Externalizar la configuracion sensible (clave secreta "
    "JWT, credenciales de base de datos) a variables de entorno, siguiendo el "
    "principio de los Twelve-Factor Apps y evitando hardcodear secretos en "
    "el codigo fuente."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                     9. REFERENCIAS BIBLIOGRAFICAS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("9. REFERENCIAS BIBLIOGRAFICAS", level=1)

add_body(
    "Las siguientes referencias fueron consultadas durante el desarrollo del "
    "proyecto y la elaboracion de este documento:",
    italic=True
)

refs = [
    "Fielding, R. T. (2000). Architectural Styles and the Design of Network-based "
    "Software Architectures (Tesis doctoral). University of California, Irvine.",

    "Hardt, D. (2012). RFC 6749 - The OAuth 2.0 Authorization Framework. "
    "Internet Engineering Task Force (IETF). https://tools.ietf.org/html/rfc6749",

    "Jones, M., Bradley, J., y Sakimura, N. (2015). RFC 7519 - JSON Web Token (JWT). "
    "Internet Engineering Task Force (IETF). https://tools.ietf.org/html/rfc7519",

    "OWASP Foundation. (2023). OWASP API Security Top 10 - 2023. "
    "https://owasp.org/www-project-api-security/",

    "Provos, N. y Mazieres, D. (1999). A Future-Adaptable Password Scheme. "
    "Proceedings of the FREENIX Track: USENIX Annual Technical Conference.",

    "Python Software Foundation. (2024). Python 3.12 Documentation. "
    "https://docs.python.org/3/",

    "Ramirez, S. (2024). FastAPI: Modern, Fast (High-Performance), Web Framework "
    "for Building APIs with Python 3.7+. https://fastapi.tiangolo.com/",

    "Samuel Colvin et al. (2024). Pydantic: Data Validation Using Python Type "
    "Annotations. https://docs.pydantic.dev/",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#                             10. ANEXOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("10. ANEXOS", level=1)

add_heading_styled("Anexo A. Archivo de Dependencias (requirements.txt)", level=2)
add_code_block(
    "fastapi\n"
    "uvicorn\n"
    "pydantic\n"
    "python-jose[cryptography]\n"
    "passlib[bcrypt]\n"
    "python-multipart"
)

add_heading_styled("Anexo B. Instrucciones de Instalacion y Ejecucion", level=2)

add_body("Paso 1: Clonar o descargar el proyecto en un directorio local.", bold=True)

add_body("Paso 2: Crear y activar un entorno virtual (recomendado):", bold=True)
add_code_block(
    "python -m venv venv\n"
    "# En Linux/macOS:\n"
    "source venv/bin/activate\n"
    "# En Windows:\n"
    "venv\\Scripts\\activate"
)

add_body("Paso 3: Instalar las dependencias:", bold=True)
add_code_block("pip install -r requirements.txt")

add_body("Paso 4: Iniciar el servidor de desarrollo:", bold=True)
add_code_block("uvicorn main:app --reload --port 8002")

add_body("Paso 5: Acceder a la documentacion interactiva:", bold=True)
add_bullet("Swagger UI: http://127.0.0.1:8002/docs")
add_bullet("ReDoc: http://127.0.0.1:8002/redoc")

add_heading_styled("Anexo C. Credenciales de Prueba", level=2)
create_table(
    headers=["Campo", "Valor"],
    data=[
        ("Usuario", "chef"),
        ("Contrasena", "cocina2024"),
        ("Rol", "Administrador / Chef"),
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#                             GUARDAR
# ═══════════════════════════════════════════════════════════════════════════════
output_path = "/home/humberto/source/home_work/proyecto_final/Proyecto_Final_API_Restaurante.docx"
doc.save(output_path)
print(f"Documento generado exitosamente: {output_path}")
